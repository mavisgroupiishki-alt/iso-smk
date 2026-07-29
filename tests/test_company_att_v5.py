"""Regression test for deterministic import of completed Forms 1–5.

Run against the user's five original DOCX files:
    SOURCE_FORMS_DIR=/mnt/data python tests/test_company_att_v5.py
"""
from __future__ import annotations

import os
from pathlib import Path
import sys
from tempfile import TemporaryDirectory

from docx import Document

BASE = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(BASE))

from company_attestation_source_parser import (  # noqa: E402
    merge_company_attestation_sources,
    parse_company_attestation_docx,
)
from generator_company_att_templates import generate_company_attestation_package_v2  # noqa: E402


SOURCE_DIR = Path(os.environ.get("SOURCE_FORMS_DIR", "/mnt/data"))
SOURCE_FILES = [
    "1. Заявление.docx",
    "2. ИТР.docx",
    "3. Трудовые.docx",
    "4. Дипломы.docx",
    "5. Аттестаты.docx",
]


def doc_text(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def by_fio(people: list[dict]) -> dict[str, dict]:
    return {str(person.get("fio", "")).strip(): person for person in people}


def main() -> None:
    paths = [SOURCE_DIR / name for name in SOURCE_FILES]
    missing = [str(path) for path in paths if not path.exists()]
    if missing:
        raise FileNotFoundError("Не найдены тестовые исходники: " + ", ".join(missing))

    fragments = [parse_company_attestation_docx(path.read_bytes(), path.name) for path in paths]
    assert all(fragments), "Одна из заполненных форм не распознана"

    merged = merge_company_attestation_sources(fragments)
    company = merged["company"]
    att = merged["company_attestation"]
    people = by_fio(att["itr"])

    assert company["name"] == "АК СтройФемили"
    assert company["director_fio"] == "Крот Александр Евгеньевич"
    assert set(people) == {
        "Крот Александр Евгеньевич",
        "Горбунов Олег Васильевич",
        "Крот Евгений Васильевич",
    }

    director = people["Крот Александр Евгеньевич"]
    assert director["stage_years"] == "3 года"
    assert director["stage_years_here"] == "2 года"
    assert director["trudovye_numbers"] == ["ПК № 2282231"]

    engineer = people["Горбунов Олег Васильевич"]
    assert engineer["stage_years"] == "49 лет"
    assert engineer["stage_years_here"] == "Менее года"
    assert engineer["trudovye_numbers"] == [
        "Трудовая книжка б/н",
        "Вкладыш б/н",
        "Вкладыш ПК № 00154290",
        "Вкладыш АТ-III № 2528394",
    ]
    assert engineer["attestat_number"] == "СТ №251921"
    assert "26.06.2031" in " ".join(engineer["attestat_form5_text"])

    foreman = people["Крот Евгений Васильевич"]
    assert foreman["stage_years"] == "28 лет"
    assert foreman["stage_years_here"] == "1 год"
    assert foreman["trudovye_numbers"] == ["ВТ-I № 2206886"]
    assert foreman["attestat_number"] == "СТ №241674"
    assert "28.11.2030" in " ".join(foreman["attestat_form5_text"])

    assert len(att["workers"]) == 16
    assert "7.6.5" in att["work_items"]

    with TemporaryDirectory() as tmp_dir:
        result = generate_company_attestation_package_v2(company, att)
        assert not result.get("warnings"), result.get("warnings")
        out_dir = Path(tmp_dir)
        generated: dict[str, str] = {}
        for item in result["docs"]:
            path = out_dir / item["name"]
            path.write_bytes(item["bytes"])
            generated[item["name"]] = doc_text(path)

        form2 = next(text for name, text in generated.items() if "2. Форма" in name)
        form3 = next(text for name, text in generated.items() if "3. Форма" in name)
        form5 = next(text for name, text in generated.items() if "5. Форма" in name)

        for token in ["3 года", "2 года", "49 лет", "Менее года", "28 лет", "1 год"]:
            assert token in form2, f"В Форме №2 отсутствует стаж: {token}"

        for token in [
            "ПК № 2282231",
            "Трудовая книжка б/н",
            "Вкладыш б/н",
            "Вкладыш ПК № 00154290",
            "Вкладыш АТ-III № 2528394",
            "ВТ-I № 2206886",
        ]:
            assert token in form2 or token in form3, f"Не перенесён документ: {token}"

        assert "Горбунов Олег Васильевич" in form5
        assert "Крот Евгений Васильевич" in form5
        assert "СТ №251921" in form5
        assert "СТ №241674" in form5

    print("V5 REGRESSION TEST PASSED")


if __name__ == "__main__":
    main()
