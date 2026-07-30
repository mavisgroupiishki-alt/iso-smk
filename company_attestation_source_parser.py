"""Deterministic parser for already completed company-attestation forms 1–5.

The application previously sent flattened DOCX text to the language model and relied on it
for copying tables. That was not reliable: experience, labour-book inserts and one of the
attestations could disappear. This module reads the Word tables directly and returns the
same JSON schema used by the generator.
"""
from __future__ import annotations

from copy import deepcopy
from io import BytesIO
import re
from typing import Any, Iterable

try:
    from docx import Document
except Exception:  # pragma: no cover - Render requirements contain python-docx
    Document = None


_DASHES = {"", "-", "—", "–"}
_RAZRYAD = ["II", "III", "IV", "V", "VI"]


def _clean(value: Any) -> str:
    text = str(value or "").replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def _lines(value: Any) -> list[str]:
    if isinstance(value, list):
        raw: Iterable[Any] = value
    else:
        raw = re.split(r"[\r\n]+", str(value or ""))
    result = []
    for item in raw:
        line = _clean(item)
        if line and line not in _DASHES:
            result.append(line)
    return result


def _cell_lines(cell) -> list[str]:
    result: list[str] = []
    for paragraph in cell.paragraphs:
        for line in re.split(r"[\r\n]+", paragraph.text or ""):
            clean = _clean(line)
            if clean:
                result.append(clean)
    if not result:
        result = _lines(cell.text)
    return result


def _full_text(doc) -> str:
    parts = [_clean(p.text) for p in doc.paragraphs if _clean(p.text)]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(_clean(c.text) for c in row.cells if _clean(c.text))
    return "\n".join(parts)


def _norm_fio(value: Any) -> str:
    text = _clean(value).lower().replace("ё", "е")
    text = re.sub(r"[^а-яa-z0-9]+", " ", text)
    return " ".join(text.split())


def _unique(values: Iterable[Any]) -> list[Any]:
    out = []
    seen = set()
    for value in values:
        key = repr(value) if isinstance(value, dict) else _clean(value).lower()
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(value)
    return out


def _nonempty(value: Any) -> bool:
    if value is None:
        return False
    if isinstance(value, str):
        return bool(_clean(value)) and _clean(value) not in _DASHES
    if isinstance(value, (list, dict)):
        return bool(value)
    return True


def _extract_number(text: str, marker: str = "") -> str:
    source = _clean(text)
    if marker == "attestat":
        m = re.search(r"\b(?:СТ|ГС)\s*№?\s*[A-Za-zА-Яа-я0-9/-]+", source, re.I)
        return _clean(m.group(0)) if m else ""
    if marker == "diploma":
        m = re.search(r"(?:Диплом\s*)?([A-ZА-ЯЁ]{1,5}(?:-[IVX]+)?\s*№?\s*\d{4,})", source, re.I)
        return _clean(m.group(1)) if m else ""
    return ""


def _extract_dates(text: str) -> list[str]:
    return re.findall(r"\b\d{2}\.\d{2}\.\d{4}\b", str(text or ""))


def _parse_diplomas(lines: list[str]) -> list[dict]:
    """Parse repeated education blocks while preserving the exact full text separately."""
    diplomas: list[dict] = []
    starts = [i for i, line in enumerate(lines) if re.search(r"\bДиплом\b", line, re.I)]
    for pos, idx in enumerate(starts):
        end = starts[pos + 1] if pos + 1 < len(starts) else len(lines)
        block = lines[idx:end]
        header = block[0]
        number = _extract_number(header, "diploma")
        dates = _extract_dates(header)
        # Usually institution/speciality/qualification follow the diploma line.
        tail = [x for x in block[1:] if x]
        education_level = lines[idx - 1] if idx > 0 and re.search(r"высш|средн", lines[idx - 1], re.I) else ""
        diplomas.append({
            "number": number,
            "date": dates[0] if dates else "",
            "institution": tail[0] if len(tail) > 0 else "",
            "speciality": tail[1] if len(tail) > 1 else "",
            "qualification": tail[2] if len(tail) > 2 else "",
            "education_level": education_level,
        })
    return [d for d in diplomas if any(_nonempty(v) for v in d.values())]


def _parse_order(lines: list[str]) -> tuple[str, str]:
    joined = " ".join(lines)
    m = re.search(r"Приказ\s*№\s*([^\s]+)(?:\s+от\s+(\d{2}\.\d{2}\.\d{4}))?", joined, re.I)
    return (_clean(m.group(1)), _clean(m.group(2))) if m else ("", "")


def _parse_attestation_lines(lines: list[str]) -> dict:
    valid = [x for x in lines if x not in _DASHES]
    if not valid:
        return {}
    joined = " ".join(valid)
    number = _extract_number(joined, "attestat")
    dates = _extract_dates(joined)
    specialization = ""
    for line in valid[1:]:
        if not re.search(r"\b(?:СТ|ГС)\b|\d{2}\.\d{2}\.\d{4}", line, re.I):
            specialization = line
            break
    if not specialization and len(valid) > 1:
        specialization = valid[-1]
    return {
        "attestat_number": number,
        "attestat_date_from": dates[0] if dates else "",
        "attestat_date_to": dates[1] if len(dates) > 1 else "",
        "attestat_specialization": specialization,
    }


def _detect_form(doc, filename: str = "") -> str | None:
    text = (_full_text(doc) + "\n" + filename).lower().replace("ё", "е")
    # Check the statement first: its attachment list mentions Forms 2–5, which must not
    # make it look like Form 2.
    if "заявление" in text and "о получении аттестата соответствия" in text:
        return "statement"
    if "форма № 2" in text or "форма №2" in text:
        return "form2"
    if "форма № 3" in text or "форма №3" in text:
        return "form3"
    if "форма № 4" in text or "форма №4" in text:
        return "form4"
    if "форма № 5" in text or "форма №5" in text:
        return "form5"
    return None


def _parse_company_statement(doc) -> dict:
    paragraphs = [_clean(p.text) for p in doc.paragraphs]
    nonempty = [x for x in paragraphs if x]
    company: dict[str, Any] = {}
    if nonempty:
        first = nonempty[0]
        m = re.search(r"^(Общество с ограниченной ответственностью|Закрытое акционерное общество|Открытое акционерное общество|Унитарное предприятие|ООО|ЗАО|ОАО|УП)\s*[\"«]?\s*(.*?)\s*[\"»]?\s*$", first, re.I)
        if m:
            long_form, name = _clean(m.group(1)), _clean(m.group(2)).strip('"«» ')
            form_map = {
                "общество с ограниченной ответственностью": "ООО",
                "закрытое акционерное общество": "ЗАО",
                "открытое акционерное общество": "ОАО",
                "унитарное предприятие": "УП",
            }
            company["form"] = form_map.get(long_form.lower(), long_form.upper())
            company["name"] = name
    for line in nonempty[:15]:
        lower = line.lower()
        if re.search(r"\bунп\b", lower):
            m = re.search(r"\bУНП\s*[:№]?\s*(\d{7,12})", line, re.I)
            if m: company["unp"] = m.group(1)
        elif re.search(r"(?:тел\.?|телефон)\s*[:.]", lower):
            m = re.search(r"(?:Тел\.?|Телефон)\s*[:.]?\s*(.+)$", line, re.I)
            if m: company["phone"] = _clean(m.group(1))
        elif "e-mail" in lower or "email" in lower or "@" in line:
            m = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-zА-Яа-я]{2,}", line.replace(" ", ""))
            if m: company["email"] = m.group(0)
        elif re.search(r"\bр/с\b|\bрасчетн", lower):
            m = re.search(r"(BY[A-Z0-9]{20,32})", line.replace(" ", ""), re.I)
            if m: company["bank_account"] = m.group(1).upper()
        elif "бик" in lower:
            m = re.search(r"БИК\s*([A-Z0-9]{8,12})", line, re.I)
            if m: company["bik"] = m.group(1).upper()
            bank = re.sub(r",?\s*БИК.*$", "", line, flags=re.I)
            bank = re.sub(r"^в\s+", "", bank, flags=re.I)
            if bank: company["bank_name"] = _clean(bank)
        elif re.match(r"^\d{6},", line) or ("область" in lower and ("г." in lower or "город" in lower)):
            company.setdefault("address", line)
    if company.get("bank_account") or company.get("bank_name") or company.get("bik"):
        bits = []
        if company.get("bank_account"): bits.append("р/с: " + company["bank_account"])
        bank_line = company.get("bank_name", "")
        if company.get("bik"): bank_line = (bank_line + ", БИК " + company["bik"]).strip(", ")
        if bank_line: bits.append(bank_line)
        company["bank_details"] = "\n".join(bits)

    full = _full_text(doc)
    # Signature is a stronger source for current director than a random mention in text.
    m = re.search(r"(?:Директор|Руководитель)\s+(?:ООО|ОАО|ЗАО|УП).*?_{3,}\s*([А-ЯЁA-Z]\.[А-ЯЁA-Z]\.\s*[А-ЯЁA-Z][а-яёa-z-]+)", full, re.S)
    if m:
        initials_surname = _clean(m.group(1))
        company["director_initials_surname"] = initials_surname
        company["director_position"] = "Директор"

    work_items = []
    in_scope = False
    for line in nonempty:
        if "на право осуществления" in line.lower():
            in_scope = True
            continue
        if in_scope and line.lower().startswith("соответствующ"):
            break
        if in_scope:
            m = re.match(r"^(7\.\d+(?:\.\d+)?)\b", line)
            if m:
                work_items.append(m.group(1))
    return {
        "company": company,
        "certification": {"standard": "company_att"},
        "company_attestation": {
            "work_items": _unique(work_items),
            "work_items_source": "document_exact",
        },
        "_source_forms": ["statement"],
    }


def _parse_form2(doc) -> dict:
    data: dict[str, Any] = {
        "certification": {"standard": "company_att"},
        "company_attestation": {"itr": [], "workers": [], "work_items_source": "document_exact"},
        "_source_forms": ["form2"],
    }
    ca = data["company_attestation"]
    paragraph_text = "\n".join(_clean(p.text) for p in doc.paragraphs if _clean(p.text))
    m = re.search(r"Общая численность работающих\s+(\d+)\s+чел", paragraph_text, re.I)
    if m: ca["staff_total"] = int(m.group(1))
    m = re.search(r"по состоянию на\s*([\d .]{8,14})", paragraph_text, re.I)
    if m:
        date = re.sub(r"\s+", "", m.group(1))
        dm = re.search(r"\d{2}\.\d{2}\.\d{4}", date)
        if dm: ca["as_of_date"] = dm.group(0)

    itr_table = next((t for t in doc.tables if len(t.columns) >= 7 and "Наименование должности" in t.rows[0].cells[1].text), None)
    if itr_table is not None:
        for row in itr_table.rows[2:]:
            cells = row.cells
            if not _clean(cells[0].text).isdigit():
                continue
            education = _cell_lines(cells[3])
            stage = _cell_lines(cells[4])
            stage = [x for x in stage if x not in _DASHES]
            labour = _cell_lines(cells[5])
            att_lines = _cell_lines(cells[6])
            order_number, hire_date = _parse_order(labour)
            labour_docs = [x for x in labour if not re.search(r"^Приказ\b", x, re.I)]
            attest = _parse_attestation_lines(att_lines)
            diplomas = _parse_diplomas(education)
            person = {
                "fio": _clean(cells[2].text),
                "position": _clean(cells[1].text),
                "education_full_text": education,
                "diplomas": diplomas,
                "diploma_number": diplomas[0].get("number", "") if diplomas else _extract_number(" ".join(education), "diploma"),
                # The completed Form 2 is only a reference. The live values are
                # calculated independently from labour-book employment periods.
                "stage_reference_total": stage[0] if len(stage) > 0 else "",
                "stage_reference_here": stage[1] if len(stage) > 1 else "",
                "stage_form2_text": stage,
                "stage_source": "document_reference" if stage else "",
                "stage_is_final": False,
                "trudovye_numbers": labour_docs,
                "trudovaya_number": labour_docs[0] if labour_docs else "",
                "trudovaya_form2_text": labour,
                "order_number": order_number,
                "hire_date": hire_date,
                "attestat_form2_text": [x for x in att_lines if x not in _DASHES],
                **attest,
                "source": "document",
            }
            ca["itr"].append({k: v for k, v in person.items() if _nonempty(v)})

    director = next((p for p in ca["itr"] if "директор" in str(p.get("position", "")).lower()
                     and "замест" not in str(p.get("position", "")).lower()), None)
    if director:
        data["company"] = {
            "director_fio": director.get("fio", ""),
            "director_position": director.get("position", "Директор"),
        }

    worker_table = next((t for t in doc.tables if len(t.columns) >= 8 and "Наименование профессий рабочих" in t.rows[0].cells[1].text), None)
    if worker_table is not None:
        for row in worker_table.rows[3:]:
            cells = row.cells
            if not _clean(cells[0].text).isdigit():
                continue
            profession = _clean(cells[1].text)
            if not profession or "итого" in profession.lower():
                continue
            for idx, grade in enumerate(_RAZRYAD, start=2):
                raw_count = re.sub(r"\D", "", _clean(cells[idx].text))
                if raw_count:
                    count = int(raw_count)
                    if count > 0:
                        ca["workers"].append({
                            "profession": profession,
                            "razryad": grade,
                            "count": count,
                            "source": "document",
                        })
    return data


def _parse_form3(doc) -> dict:
    ca = {"itr": []}
    table = next((t for t in doc.tables if len(t.columns) >= 4 and "Номер трудовой книжки" in t.rows[0].cells[3].text), None)
    if table is not None:
        for row in table.rows[2:]:
            if not _clean(row.cells[0].text).isdigit():
                continue
            docs = [x for x in _cell_lines(row.cells[3]) if x not in _DASHES]
            ca["itr"].append({
                "fio": _clean(row.cells[1].text),
                "position": _clean(row.cells[2].text),
                "trudovye_numbers": docs,
                "trudovaya_number": docs[0] if docs else "",
                "source": "document",
            })
    return {"certification": {"standard": "company_att"}, "company_attestation": ca, "_source_forms": ["form3"]}


def _parse_form4(doc) -> dict:
    ca = {"itr": []}
    table = next((t for t in doc.tables if len(t.columns) >= 3 and "Номер диплома" in t.rows[0].cells[2].text), None)
    if table is not None:
        for row in table.rows[2:]:
            if not _clean(row.cells[0].text).isdigit():
                continue
            numbers = [x for x in _cell_lines(row.cells[2]) if x not in _DASHES]
            ca["itr"].append({
                "fio": _clean(row.cells[1].text),
                "diploma_numbers": numbers,
                "diploma_number": numbers[0] if numbers else "",
                "source": "document",
            })
    return {"certification": {"standard": "company_att"}, "company_attestation": ca, "_source_forms": ["form4"]}


def _parse_form5(doc) -> dict:
    ca = {"itr": []}
    table = next((t for t in doc.tables if len(t.columns) >= 4 and "квалификационного аттестата" in t.rows[0].cells[3].text.lower()), None)
    if table is not None:
        for row in table.rows[2:]:
            if not _clean(row.cells[0].text).isdigit():
                continue
            att_lines = [x for x in _cell_lines(row.cells[3]) if x not in _DASHES]
            attest = _parse_attestation_lines(att_lines)
            ca["itr"].append({
                "fio": _clean(row.cells[1].text),
                "position": _clean(row.cells[2].text),
                "attestat_form5_text": att_lines,
                **attest,
                "source": "document",
            })
    return {"certification": {"standard": "company_att"}, "company_attestation": ca, "_source_forms": ["form5"]}


def parse_company_attestation_docx(file_bytes: bytes, filename: str = "") -> dict | None:
    """Return structured data when a DOCX is one of the completed forms 1–5."""
    if not file_bytes or Document is None or not filename.lower().endswith(".docx"):
        return None
    try:
        doc = Document(BytesIO(file_bytes))
    except Exception:
        return None
    form = _detect_form(doc, filename)
    if form == "statement":
        return _parse_company_statement(doc)
    if form == "form2":
        return _parse_form2(doc)
    if form == "form3":
        return _parse_form3(doc)
    if form == "form4":
        return _parse_form4(doc)
    if form == "form5":
        return _parse_form5(doc)
    return None


def _merge_person(base: dict, patch: dict, source_wins: bool = True) -> dict:
    out = deepcopy(base or {})
    list_fields = {
        "education_full_text", "diplomas", "diploma_numbers", "trudovye_numbers",
        "stage_form2_text", "trudovaya_form2_text", "attestat_form2_text", "attestat_form5_text",
        "employment_periods", "attestats", "stage_review_reasons",
    }
    for key, value in (patch or {}).items():
        if not _nonempty(value):
            continue
        if key in list_fields:
            existing = out.get(key) or []
            if source_wins and key in {
                "education_full_text", "diploma_numbers", "trudovye_numbers", "stage_form2_text",
                "trudovaya_form2_text", "attestat_form2_text", "attestat_form5_text",
            }:
                out[key] = deepcopy(value)
            else:
                out[key] = _unique([*existing, *deepcopy(value)])
        elif source_wins or not _nonempty(out.get(key)):
            out[key] = deepcopy(value)
    # Derive the single-number compatibility fields used by the current generator.
    if not _nonempty(out.get("trudovaya_number")) and out.get("trudovye_numbers"):
        out["trudovaya_number"] = out["trudovye_numbers"][0]
    if not _nonempty(out.get("diploma_number")):
        if out.get("diploma_numbers"):
            out["diploma_number"] = out["diploma_numbers"][0]
        elif out.get("diplomas"):
            out["diploma_number"] = out["diplomas"][0].get("number", "")
    if not _nonempty(out.get("attestat_number")):
        lines = out.get("attestat_form5_text") or out.get("attestat_form2_text") or []
        derived = _extract_number(" ".join(lines), "attestat")
        if derived:
            out["attestat_number"] = derived
    return out


def merge_itr_records(records: Iterable[dict]) -> list[dict]:
    """Merge duplicate people by FIO without dropping fields from other forms."""
    order: list[str] = []
    merged: dict[str, dict] = {}
    for record in records or []:
        if not isinstance(record, dict):
            continue
        key = _norm_fio(record.get("fio")) or f"__row_{len(order)}"
        if key not in merged:
            order.append(key)
            merged[key] = {}
        merged[key] = _merge_person(merged[key], record, source_wins=True)
    return [merged[key] for key in order]


def merge_company_attestation_sources(fragments: Iterable[dict]) -> dict:
    """Combine statement + Forms 2–5 into one authoritative frontend data object."""
    out: dict[str, Any] = {}
    people: list[dict] = []
    forms: list[str] = []
    for fragment in fragments or []:
        if not isinstance(fragment, dict):
            continue
        forms.extend(fragment.get("_source_forms") or [])
        company = fragment.get("company") or {}
        if company:
            out.setdefault("company", {})
            for key, value in company.items():
                if _nonempty(value):
                    out["company"][key] = deepcopy(value)
        cert = fragment.get("certification") or {}
        if cert:
            out.setdefault("certification", {}).update({k: deepcopy(v) for k, v in cert.items() if _nonempty(v)})
        ca = fragment.get("company_attestation") or {}
        if ca:
            target = out.setdefault("company_attestation", {})
            for key, value in ca.items():
                if key == "itr":
                    people.extend(value or [])
                elif key in ("work_items", "workers"):
                    if value:
                        # Form 2 workers and statement work items are authoritative exact lists.
                        target[key] = deepcopy(value)
                elif _nonempty(value):
                    target[key] = deepcopy(value)
    if people:
        out.setdefault("company_attestation", {})["itr"] = merge_itr_records(people)
    if forms:
        out["_source_forms"] = _unique(forms)
    if out.get("company_attestation"):
        out.setdefault("certification", {})["standard"] = "company_att"
    return out


def is_completed_attestation_form_text(block: str) -> bool:
    """Used by archive reconciliation so root Forms 2–5 are not collapsed into company details."""
    text = str(block or "").lower().replace("ё", "е")
    return any(
        marker in text
        for marker in (
            "форма № 2", "форма №2", "форма № 3", "форма №3",
            "форма № 4", "форма №4", "форма № 5", "форма №5",
        )
    ) or ("заявление" in text and "о получении аттестата соответствия" in text)
