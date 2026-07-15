"""
Модуль генерации документов на аттестацию ЮРИДИЧЕСКОГО ЛИЦА (компании) —
аттестат соответствия (СТ — подряд, ГС — генподряд).

АРХИТЕКТУРА: документы строятся программно через python-docx (реальные таблицы Word,
не текстовые заглушки), без обращения к ИИ для текста самого документа — только
подстановка данных в фиксированную структуру. Структура и формулировки выверены
построчно по реальным поданным и принятым документам (ООО «Асецкий и К» — без
генподряда, ЧУП «СК76» — с генподрядом). Это даёт гарантированно точный формат
вместо непредсказуемого результата ИИ-генерации текста.
"""
import json, re, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

BASE_DIR = Path(__file__).parent.resolve()
_CLASSIFIER_PATH = BASE_DIR / 'classifier_company_att.json'
if not _CLASSIFIER_PATH.exists():
    raise FileNotFoundError(
        f"Не найден classifier_company_att.json по пути {_CLASSIFIER_PATH}. "
        f"Файл должен лежать в той же папке репозитория, что и server.py/generator.py."
    )
CLASSIFIER = json.loads(_CLASSIFIER_PATH.read_text('utf-8'))

# ── Склонение организационно-правовой формы (именительный/родительный/дательный) ──
# Небольшой фиксированный список — все формы, которые реально встречаются у клиентов.
LEGAL_FORMS = {
    'ООО':  {'nom': 'Общество с ограниченной ответственностью',
             'gen': 'Общества с ограниченной ответственностью',
             'dat': 'Обществу с ограниченной ответственностью', 'quote': '«»'},
    'ОДО':  {'nom': 'Общество с дополнительной ответственностью',
             'gen': 'Общества с дополнительной ответственностью',
             'dat': 'Обществу с дополнительной ответственностью', 'quote': '«»'},
    'ЗАО':  {'nom': 'Закрытое акционерное общество',
             'gen': 'Закрытого акционерного общества',
             'dat': 'Закрытому акционерному обществу', 'quote': '«»'},
    'ОАО':  {'nom': 'Открытое акционерное общество',
             'gen': 'Открытого акционерного общества',
             'dat': 'Открытому акционерному обществу', 'quote': '«»'},
    'ЧУП':  {'nom': 'Частное унитарное предприятие',
             'gen': 'Частного унитарного предприятия',
             'dat': 'Частному унитарному предприятию', 'quote': '""'},
    'ЧТУП': {'nom': 'Частное торговое унитарное предприятие',
             'gen': 'Частного торгового унитарного предприятия',
             'dat': 'Частному торговому унитарному предприятию', 'quote': '""'},
    'ИП':   {'nom': 'Индивидуальный предприниматель',
             'gen': 'Индивидуального предпринимателя',
             'dat': 'Индивидуальному предпринимателю', 'quote': '""'},
}


def _legal(form):
    return LEGAL_FORMS.get((form or 'ООО').upper(), LEGAL_FORMS['ООО'])


def _quoted_name(company, case='nom'):
    """Полное название с формой собственности в нужном падеже и правильными кавычками."""
    L = _legal(company.get('form'))
    name = company.get('name', '')
    q = L['quote']
    return f"{L[case]} {q[0]}{name}{q[1]}"


def _normalize_category(category):
    """Игорь иногда пишет в JSON текстовую строку "null" вместо настоящего null.
    Приводим все "пустые" варианты к настоящему None."""
    if category is None:
        return None
    s = str(category).strip().lower()
    if s in ('', 'null', 'none', 'нет', 'undefined'):
        return None
    return str(category).strip()


def find_work_items(query: str, max_items=10):
    """Ищет пункты 7.x классификатора по свободному тексту клиента (по началу слова —
    русский язык сильно склоняется)."""
    q = query.lower()
    q_stems = {w[:5] for w in re.findall(r'[а-яё]{5,}', q)}
    items = CLASSIFIER['punkt_7_smr']['items']
    found = []
    for code, text in items.items():
        tl = text.lower()
        t_stems = re.findall(r'[а-яё]{5,}', tl)
        score = sum(1 for w in t_stems if w[:5] in q_stems)
        if score > 0:
            found.append((score, code, text))
    found.sort(key=lambda x: -x[0])
    return [(code, text) for _, code, text in found[:max_items]]


def check_category_requirements(category, staff_total: int, has_smetchik: bool,
                                 experience_objects: list, prior_category_years: int = 0) -> list:
    """Возвращает список предупреждений (пустой список = всё ок)."""
    category = _normalize_category(category)
    if category is None:
        return []
    warnings = []
    thresholds = CLASSIFIER['_meta']['category_thresholds'].get(str(category))
    if not thresholds:
        return [f"Категория '{category}' не входит в список 1-4 — проверьте, что имелось в виду."]
    if staff_total < thresholds['min_staff']:
        warnings.append(
            f"Недостаточно штата для категории {category}: нужно минимум {thresholds['min_staff']} чел. "
            f"по основному месту работы, у клиента {staff_total}. Реалистичный вариант — категория ниже."
        )
    if not has_smetchik:
        warnings.append("Нет аттестованного инженера по сметной работе (сметчика) — обязателен для любой категории генподряда.")
    objects_required = thresholds['objects_required']
    if objects_required > 0:
        n_objects = len(experience_objects or [])
        if n_objects < objects_required:
            warnings.append(
                f"Для категории {category} нужно подтвердить опыт минимум по {objects_required} объектам "
                f"(генподрядчик, привлекавший субподряд, введён в эксплуатацию не позднее 5 лет назад, "
                f"не текущий ремонт). Предоставлено объектов: {n_objects}."
            )
        if thresholds.get('prior_category') and prior_category_years < thresholds.get('prior_years', 0):
            warnings.append(
                f"Для категории {category} нужно не менее {thresholds['prior_years']} лет владения категорией "
                f"{thresholds['prior_category']} — по данным клиента стаж владения {prior_category_years} лет."
            )
    return warnings


def calculate_stazh(periods: list, as_of_date: str = None) -> dict:
    """
    Считает трудовой стаж по календарному методу — та же логика, что использует
    типовой онлайн-калькулятор стажа (на основе Приказа Минздравсоцразвития РФ
    от 06.02.2007 № 91): каждые 30 дней = 1 месяц, каждые 12 месяцев = 1 год.

    periods: [{"start": "ДД.ММ.ГГГГ", "end": "ДД.ММ.ГГГГ" или None (по настоящее время)}, ...]
              — берутся буквально из записей трудовой книжки (даты приёма/увольнения).
    as_of_date: на какую дату считать (по умолчанию — сегодня).

    Возвращает {'years': int, 'months': int, 'days': int, 'total_years_rounded': float}.
    Игорю НЕ нужно самому считать стаж в уме — он просто переписывает даты приёма/увольнения
    из трудовой книжки в periods[], а точный расчёт делает эта функция.
    """
    from datetime import datetime as _dt

    def _parse(d):
        if not d:
            return None
        for fmt in ('%d.%m.%Y', '%d.%m.%y', '%Y-%m-%d'):
            try:
                return _dt.strptime(d.strip(), fmt)
            except (ValueError, AttributeError):
                continue
        return None

    today = _parse(as_of_date) or _dt.now()
    total_days = 0
    for period in (periods or []):
        start = _parse(period.get('start'))
        end = _parse(period.get('end')) or today
        if not start:
            continue
        if end < start:
            continue
        total_days += (end - start).days + 1  # включительно, как в трудовом законодательстве

    months, days = divmod(total_days, 30)
    years, months = divmod(months, 12)

    return {
        'years': years,
        'months': months,
        'days': days,
        'total_years_rounded': round(years + months / 12 + days / 365, 1),
        'display': f"{years} лет {months} мес. {days} дн." if years or months else f"{days} дн.",
    }


def _dir_init(fio: str) -> str:
    parts = (fio or '').strip().split()
    if len(parts) >= 3:
        return f"{parts[1][0]}.{parts[2][0]}. {parts[0]}"
    return fio or ''


# ── Низкоуровневые помощники построения документа ──────────────────────────
def _new_doc(landscape=False):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(2.5) if not landscape else Cm(1.5)
    sec.right_margin = Cm(1.5)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
    return doc


def _p(doc, text='', align=None, bold=False, size=11, space_after=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
    return para


def _table(doc, headers, rows, widths=None):
    n = len(headers)
    t = doc.add_table(rows=1, cols=n)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val) if val not in (None, '') else '—')
            run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return t


def _doc_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Документ 1: Заявление ───────────────────────────────────────────────────
def gen_zayavlenie_company(company: dict, work_items: list, category: str) -> bytes:
    category = _normalize_category(category)
    L = _legal(company.get('form'))
    full_nom = _quoted_name(company, 'nom')
    full_gen = _quoted_name(company, 'gen')
    full_dat = _quoted_name(company, 'dat')
    dir_pos = company.get('director_position', 'Директор')
    dir_init = _dir_init(company.get('director_fio', ''))

    doc = _new_doc()

    # Шапка заявителя
    _p(doc, full_nom, bold=True)
    _p(doc, company.get('address', ''))
    _p(doc, f"р/с: {company.get('bank_details','')}" if company.get('bank_details') else '')
    _p(doc, f"УНП {company.get('unp','')}")
    _p(doc, f"Тел./факс: {company.get('phone','')}")
    _p(doc, f"e-mail: {company.get('email','')}")
    _p(doc, "")
    _p(doc, "Исх. № ___ от ___.___.____ г.", align='right')
    _p(doc, "")
    _p(doc, "РУП «БЕЛСТРОЙЦЕНТР»")
    _p(doc, "ул. Р. Люксембург, 101")
    _p(doc, "220036, г. Минск")
    _p(doc, "")
    _p(doc, full_nom)
    _p(doc, company.get('address', ''))
    _p(doc, f"УНП {company.get('unp','')}")
    _p(doc, f"Тел.: {company.get('phone','')}")
    _p(doc, f"e-mail: {company.get('email','')}")
    _p(doc, "")

    _p(doc, "ЗАЯВЛЕНИЕ", align='center', bold=True, size=13)
    _p(doc, "о получении аттестата соответствия", align='center', bold=True)
    _p(doc, "")

    _p(doc, f"Прошу провести аттестацию {full_gen} на право осуществления:", align='justify')

    if category:
        _p(doc, f"6. Выполнение функций генерального подрядчика со стоимостью строительства свыше "
                f"{CLASSIFIER['_meta']['genpodryad_min_cost']}. Соответствующей квалификационным "
                f"требованиям, предъявляемым для получения аттестата соответствия {category} "
                f"класса(ов) сложности.", align='justify')

    _p(doc, "7. Выполнение строительно-монтажных работ:", align='justify')
    for code in work_items:
        text = CLASSIFIER['punkt_7_smr']['items'].get(code, code)
        _p(doc, f"{code}. {text};", align='justify')

    _p(doc, "соответствующей квалификационным требованиям, предъявляемым для получения "
            "аттестатов(а) соответствия 1-4 классов(а) сложности.", align='justify')
    _p(doc, "")
    _p(doc, "Сведения об обособленных подразделениях, в том числе филиалах (при их наличии): нет")
    _p(doc, "")
    _p(doc, f"В соответствии с {CLASSIFIER['_meta']['legal_basis']} прошу оформить {full_dat} "
            f"аттестат соответствия на бумажном носителе. Сведения, изложенные в заявлении и "
            f"прилагаемых к нему документах, достоверны.", align='justify')
    _p(doc, "")

    _p(doc, "Приложение:", bold=True)
    prilozhenie_rows = [
        ["1", "Легализованная выписка из торгового реестра страны, в которой иностранная "
              "организация учреждена, или иное эквивалентное доказательство юридического статуса "
              "иностранной организации в соответствии с законодательством страны ее учреждения "
              "(для заявителя – нерезидента).", ""],
        ["2", "Сведения о составе и профессиональной квалификации руководящих работников, "
              "специалистов и рабочих, работающих по основному месту работы (форма № 2).", ""],
        ["3", "Сводный список и копии трудовых книжек руководящих работников, специалистов, "
              "работающих по основному месту работы (форма № 3).", ""],
        ["4", "Сводный список и копии дипломов руководящих работников, специалистов, работающих "
              "по основному месту работы (форма № 4).", ""],
        ["5", "Сводный список и копии квалификационных аттестатов руководящих работников, "
              "специалистов, работающих по основному месту работы (форма № 5).", ""],
    ]
    if category:
        prilozhenie_rows.append(
            ["6", "Сведения о наличии опыта выполнения работ (оказания услуг) по заявляемому виду "
                  "деятельности в области строительства за последние пять лет в качестве "
                  "генерального подрядчика (форма № 6).", ""]
        )
    _table(doc, ["№ п/п", "Наименование документа", "Кол-во листов"], prilozhenie_rows,
           widths=[1.5, 14, 2.5])
    _p(doc, "")
    _p(doc, "Всего:")
    _p(doc, "")
    _p(doc, f"{dir_pos} _____________ {dir_init}")

    return _doc_bytes(doc)


# ── Заявление на отмену ─────────────────────────────────────────────────────
def gen_zayavlenie_otmena(company: dict, old_attestat_number: str, reason: str) -> bytes:
    full_nom = _quoted_name(company, 'nom')
    dir_init = _dir_init(company.get('director_fio', ''))
    doc = _new_doc()
    _p(doc, full_nom, bold=True)
    _p(doc, company.get('address', ''))
    _p(doc, f"УНП {company.get('unp','')}")
    _p(doc, f"Тел.: {company.get('phone','')}")
    _p(doc, f"e-mail: {company.get('email','')}")
    _p(doc, "")
    _p(doc, "Исх. № ___ от ___.___.____ г.", align='right')
    _p(doc, "")
    _p(doc, "РУП «БЕЛСТРОЙЦЕНТР»")
    _p(doc, "ул. Р. Люксембург, 101")
    _p(doc, "220036, г. Минск")
    _p(doc, "")
    _p(doc, full_nom)
    _p(doc, company.get('address', ''))
    _p(doc, f"УНП {company.get('unp','')}")
    _p(doc, "")
    _p(doc, "ЗАЯВЛЕНИЕ", align='center', bold=True, size=13)
    _p(doc, "о прекращении действия аттестата соответствия", align='center', bold=True)
    _p(doc, "")
    _p(doc, f"{full_nom} просит прекратить действие выданного ранее аттестата соответствия "
            f"от ___.___.____ г. № {old_attestat_number}.", align='justify')
    _p(doc, "")
    _p(doc, f"Причина: {reason}", align='justify')
    _p(doc, "")
    _p(doc, "В соответствии со статьёй 36 Кодекса Республики Беларусь об архитектурной, "
            "градостроительной и строительной деятельности.", align='justify')
    _p(doc, "")
    _p(doc, f"Директор _____________ {dir_init}")
    return _doc_bytes(doc)


# ── Документ 2: Форма №2 — ИТР + рабочие ────────────────────────────────────
def gen_form2_itr(company: dict, itr_list: list, workers: list, work_scope_text: str) -> bytes:
    """Landscape — таблица широкая (7 колонок). Рабочие — второй раздел внутри
    ЭТОЙ ЖЕ формы (не отдельный документ), как в реальном образце ЧУП «СК76»."""
    full_nom = _quoted_name(company, 'nom')
    dir_init = _dir_init(company.get('director_fio', ''))
    total_staff = company.get('staff_total') or (len(itr_list) + sum(w.get('count', 0) or 0 for w in workers))
    n_itr = len(itr_list)

    doc = _new_doc(landscape=True)
    _p(doc, full_nom, bold=True)
    _p(doc, "")
    _p(doc, "Форма № 2", align='center', bold=True)
    _p(doc, "СВЕДЕНИЯ о составе и профессиональной квалификации руководящих работников, "
            "специалистов и рабочих, работающих по основному месту работы", align='center', bold=True)
    _p(doc, "")
    _p(doc, f"Общая численность работающих {total_staff} чел., в том числе по заявляемому виду "
            f"деятельности {total_staff} чел. по состоянию на ___.___.____ ; численность "
            f"инженерно-технических работников по заявляемому виду деятельности {n_itr} чел.")
    _p(doc, f"Область деятельности: {work_scope_text}")
    _p(doc, "")

    itr_rows = []
    for i, p_ in enumerate(itr_list, 1):
        obrazovanie = (f"{p_.get('education_level','')}, диплом {p_.get('diploma_number') or '—'} "
                       f"выдан {p_.get('diploma_date') or '—'}, {p_.get('diploma_institution','')}, "
                       f"{p_.get('diploma_speciality','')}, {p_.get('diploma_qualification','')}")
        stazh = (f"{p_.get('stage_years') or '—'} лет / {p_.get('stage_years_here') or '—'}")
        trudovaya = (f"{p_.get('trudovaya_number') or '—'}, Приказ №{p_.get('order_number') or '—'} "
                     f"от {p_.get('hire_date') or '—'}")
        attestat = p_.get('attestat_number') or '—'
        if p_.get('attestat_date'):
            attestat += f" от {p_.get('attestat_date')}"
        if p_.get('attestat_specialization'):
            attestat += f" {p_.get('attestat_specialization')}"
        itr_rows.append([i, p_.get('position',''), p_.get('fio',''), obrazovanie, stazh, trudovaya, attestat])

    _table(doc, ["№", "Должность", "ФИО", "Образование (уровень, диплом, учреждение, специальность, квалификация)",
                 "Стаж (по виду деятельности / у нанимателя)", "Трудовая книжка + приказ", "Аттестат, специализация"],
           itr_rows, widths=[1, 3, 3.5, 6, 2.5, 3.5, 4])

    _p(doc, "")
    _p(doc, "Раздел 2 — рабочие строительных профессий, соответствующих заявляемым видам "
            "деятельности в области строительства согласно технологической документации на "
            "производство строительно-монтажных работ, работающих по основному месту работы:")
    if workers:
        w_rows = [[i, w.get('profession',''), w.get('razryad','') or '—', w.get('count','') or '—']
                  for i, w in enumerate(workers, 1)]
        _table(doc, ["№", "Профессия рабочего", "Разряд", "Количество человек"], w_rows,
               widths=[1, 8, 3, 4])
    else:
        _p(doc, "Сведения о рабочих не предоставлены на момент подготовки документа.")

    _p(doc, "")
    _p(doc, f"Директор {full_nom} _____________ {dir_init}")
    _p(doc, "«___» _______ 202_ г.")
    return _doc_bytes(doc)


# ── Документы 3-5: сводные списки (простые таблицы) ─────────────────────────
def gen_form3_trudovye(company: dict, itr_list: list) -> bytes:
    full_nom = _quoted_name(company, 'nom')
    dir_init = _dir_init(company.get('director_fio', ''))
    doc = _new_doc()
    _p(doc, full_nom, bold=True)
    _p(doc, "")
    _p(doc, "Форма № 3", align='center', bold=True)
    _p(doc, "СВОДНЫЙ СПИСОК трудовых книжек руководящих работников, специалистов, работающих по "
            "основному месту работы", align='center', bold=True)
    _p(doc, "")
    rows = [[i, p_.get('fio',''), p_.get('position',''), p_.get('trudovaya_number') or '—']
            for i, p_ in enumerate(itr_list, 1)]
    _table(doc, ["№ п/п", "Ф.И.О.", "Должность в соответствии с записью в трудовой книжке",
                 "Номер трудовой книжки"], rows, widths=[1.5, 5, 7, 4])
    _p(doc, "")
    _p(doc, f"Директор _____________ {dir_init}")
    _p(doc, "«___» _______ 202_ г.")
    return _doc_bytes(doc)


def gen_form4_diplomy(company: dict, itr_list: list) -> bytes:
    full_nom = _quoted_name(company, 'nom')
    dir_init = _dir_init(company.get('director_fio', ''))
    doc = _new_doc()
    _p(doc, full_nom, bold=True)
    _p(doc, "")
    _p(doc, "Форма № 4", align='center', bold=True)
    _p(doc, "СВОДНЫЙ СПИСОК дипломов руководящих работников, специалистов, работающих по "
            "основному месту работы", align='center', bold=True)
    _p(doc, "")
    rows = [[i, p_.get('fio',''), p_.get('diploma_number') or '—'] for i, p_ in enumerate(itr_list, 1)]
    _table(doc, ["№ п/п", "Ф.И.О.", "Номер диплома"], rows, widths=[1.5, 8, 6])
    _p(doc, "")
    _p(doc, f"Директор _____________ {dir_init}")
    _p(doc, "«___» _______ 202_ г.")
    return _doc_bytes(doc)


def gen_form5_attestaty(company: dict, itr_list: list) -> bytes:
    full_nom = _quoted_name(company, 'nom')
    dir_init = _dir_init(company.get('director_fio', ''))
    doc = _new_doc()
    _p(doc, full_nom, bold=True)
    _p(doc, "")
    _p(doc, "Форма № 5", align='center', bold=True)
    _p(doc, "СВОДНЫЙ СПИСОК квалификационных аттестатов руководящих работников, специалистов, "
            "работающих по основному месту работы", align='center', bold=True)
    _p(doc, "")
    rows = []
    for i, p_ in enumerate(itr_list, 1):
        att = p_.get('attestat_number', '')
        if att:
            info = f"{att} с {p_.get('attestat_date_from','')} г. по {p_.get('attestat_date_to','')} г. {p_.get('attestat_specialization','')}"
        else:
            info = "нет аттестата / в процессе получения"
        rows.append([i, p_.get('fio',''), p_.get('position',''), info])
    _table(doc, ["№ п/п", "Ф.И.О.", "Должность", "Номер и срок действия аттестата, специализация"],
           rows, widths=[1.5, 5, 4, 7])
    _p(doc, "")
    _p(doc, f"Директор _____________ {dir_init}")
    _p(doc, "«___» _______ 202_ г.")
    return _doc_bytes(doc)


def gen_form6_opyt(company: dict, experience_objects: list) -> bytes:
    full_nom = _quoted_name(company, 'nom')
    dir_init = _dir_init(company.get('director_fio', ''))
    doc = _new_doc()
    _p(doc, full_nom, bold=True)
    _p(doc, "")
    _p(doc, "Форма № 6", align='center', bold=True)
    _p(doc, "Сведения о наличии опыта выполнения работ (оказания услуг) по заявляемому виду "
            "деятельности в области строительства за последние пять лет в качестве генерального "
            "подрядчика", align='center', bold=True)
    _p(doc, "")
    if experience_objects:
        rows = [[i, o.get('name',''), o.get('complexity_class','')] for i, o in enumerate(experience_objects, 1)]
    else:
        rows = [[1, '-', '-'], [2, '-', '-']]
    _table(doc, ["№", "Наименование объекта", "Класс сложности согласно СН 3.02.07-2020"],
           rows, widths=[1.5, 10, 5])
    _p(doc, "")
    _p(doc, f"Директор _____________ {dir_init}")
    return _doc_bytes(doc)


# ── Главный конвейер ─────────────────────────────────────────────────────────
def generate_company_attestation_package(company: dict, attestation_data: dict, api_key, vibe_call_fn,
                                          progress_cb=None) -> dict:
    """
    api_key/vibe_call_fn больше не используются для текста документов (генерация теперь
    полностью детерминирована через python-docx) — оставлены в сигнатуре для совместимости
    с существующим вызовом из generator.py.
    """
    docs = []
    step = [0]
    category_for_total = _normalize_category(attestation_data.get('category'))
    if attestation_data.get('is_cancellation'):
        total_steps = 1
    elif category_for_total:
        total_steps = 6
    else:
        total_steps = 5

    def p(msg):
        step[0] += 1
        if progress_cb:
            progress_cb(step[0], total_steps, msg)
        print(f"  [company_att {step[0]}] {msg}")

    org = company.get('name', 'company')
    category = _normalize_category(attestation_data.get('category'))
    itr_list = attestation_data.get('itr', [])
    workers = attestation_data.get('workers', [])
    staff_total = attestation_data.get('staff_total', len(itr_list))
    has_smetchik = attestation_data.get('has_smetchik', False)
    experience_objects = attestation_data.get('experience_objects', [])
    prior_years = attestation_data.get('prior_category_years', 0)

    # Автоматический расчёт стажа из сырых дат (Игорю не нужно считать в уме — только
    # переписать даты приёма/увольнения из трудовой книжки в employment_periods).
    for person in itr_list:
        periods = person.get('employment_periods')
        if periods and not person.get('stage_years'):
            calc = calculate_stazh(periods, as_of_date=attestation_data.get('as_of_date'))
            person['stage_years'] = calc['display']
        # Стаж именно у текущего нанимателя — последний период в списке (если явно не указан)
        if periods and not person.get('stage_years_here'):
            last_period = periods[-1] if periods else None
            if last_period:
                calc_here = calculate_stazh([last_period], as_of_date=attestation_data.get('as_of_date'))
                person['stage_years_here'] = calc_here['display']

    warnings = []
    if category:
        warnings = check_category_requirements(category, staff_total, has_smetchik, experience_objects, prior_years)

    if len(itr_list) <= 1 and staff_total > 1:
        warnings.append(
            f"В данных только {len(itr_list)} человек в ИТР, хотя штат указан как {staff_total} — "
            f"похоже часть людей потерялась при разборе. Проверьте пакет перед подачей."
        )
    empty_itr = [p_.get('fio', f'#{i+1}') for i, p_ in enumerate(itr_list)
                 if not p_.get('diploma_number') and not p_.get('stage_years') and not p_.get('trudovaya_number')]
    if empty_itr:
        warnings.append(
            f"У этих людей вообще не заполнены диплом/стаж/трудовая (в документе будут прочерки): "
            f"{', '.join(empty_itr)}."
        )
    partial_missing_trudovaya = [p_.get('fio', '?') for p_ in itr_list
                                  if p_.get('diploma_number') and not p_.get('trudovaya_number')]
    if partial_missing_trudovaya:
        warnings.append(
            f"У этих людей есть диплом, но нет номера трудовой книжки: {', '.join(partial_missing_trudovaya)}."
        )
    if not workers:
        warnings.append(
            "Реальные данные о рабочих не переданы — раздел «рабочие» в Форме №2 будет пустым, "
            "а не придуман по виду работ. Уточните у клиента список профессий/разрядов/количества."
        )
    elif all(not w.get('count') for w in workers):
        warnings.append("В разделе «рабочие» Формы №2 не указано количество человек по профессиям.")

    if attestation_data.get('is_cancellation'):
        p("Заявление на отмену/исключение")
        docs.append({
            'name': f"{org} - Заявление на отмену.docx",
            'bytes': gen_zayavlenie_otmena(
                company, attestation_data.get('old_attestat_number', ''),
                attestation_data.get('cancellation_reason', 'по заявлению обладателя')
            )
        })
        return {'docs': docs, 'warnings': warnings}

    work_items = attestation_data.get('work_items') or []
    if not work_items and attestation_data.get('work_scope_text'):
        found = find_work_items(attestation_data['work_scope_text'])
        work_items = [code for code, _ in found]
    if not work_items:
        work_items = ['7.4.1']

    p("1. Заявление")
    docs.append({'name': f"{org} - 1. Заявление.docx",
                  'bytes': gen_zayavlenie_company(company, work_items, category)})

    work_scope_text = ', '.join(CLASSIFIER['punkt_7_smr']['items'].get(c, c) for c in work_items)

    p("2. Форма №2 (ИТР и рабочие)")
    docs.append({'name': f"{org} - 2. Форма №2 ИТР и рабочие.docx",
                  'bytes': gen_form2_itr(company, itr_list, workers, work_scope_text)})

    p("3. Форма №3 (Трудовые)")
    docs.append({'name': f"{org} - 3. Форма №3 Трудовые.docx",
                  'bytes': gen_form3_trudovye(company, itr_list)})

    p("4. Форма №4 (Дипломы)")
    docs.append({'name': f"{org} - 4. Форма №4 Дипломы.docx",
                  'bytes': gen_form4_diplomy(company, itr_list)})

    p("5. Форма №5 (Аттестаты)")
    docs.append({'name': f"{org} - 5. Форма №5 Аттестаты.docx",
                  'bytes': gen_form5_attestaty(company, itr_list)})

    if category:
        p("6. Форма №6 (Опыт генподрядчика)")
        docs.append({'name': f"{org} - 6. Форма №6 Опыт.docx",
                      'bytes': gen_form6_opyt(company, experience_objects)})

    return {'docs': docs, 'warnings': warnings}
