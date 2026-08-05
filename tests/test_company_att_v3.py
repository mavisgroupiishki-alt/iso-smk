from pathlib import Path
import sys
from docx import Document

BASE = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(BASE))
from generator_company_att_templates import generate_company_attestation_package_v2

company = {
    'form': 'ООО',
    'name': 'АК СтройФемили',
    'address': '212027, Могилевская область, г. Могилев, ул. Гончарная, д. 3, каб. 5',
    'bank_details': 'р/с: BY94POIS30120162570701933001\nв ОАО «Паритетбанк», БИК POISBY2X',
    'unp': '791371126',
    'phone': '+375 29 120-32-66',
    'email': 'aliaksandrkrot@mail.ru',
    'director_fio': 'Крот Александр Евгеньевич',
    'director_position': 'Директор',
}
att = {
    'category': None,
    'work_scope_text': 'Общестроительные работы',
    # старый неполный результат ИИ: генератор должен расширить его до 7.2–7.6
    'work_items': ['7.4.1', '7.4.3', '7.4.4'],
    'work_items_source': 'auto',
    'as_of_date': '15.06.2026',
    'staff_total': 3,
    # старый неправильный auto-набор должен быть заменён полным набором 16 работников
    'workers': [
        {'profession':'Каменщик','razryad':'IV','count':2,'source':'auto'},
        {'profession':'Бетонщик','razryad':'IV','count':2,'source':'auto'},
        {'profession':'Монтажник строительных конструкций','razryad':'IV','count':2,'source':'auto'},
    ],
    'itr': [
        {
            'fio': 'Крот Александр Евгеньевич',
            'position': 'Директор',
            'education_full_text': [
                'Высшее',
                'Диплом А № 1554029 выдан 24.06.2022 г.',
                'Белорусская государственная орденов Октябрьской Революции и Трудового Красного Знамени сельскохозяйственная академия',
                'Сельское строительство и обустройство территорий',
                'Инженер',
            ],
            'diplomas': [{'number':'А № 1554029','date':'24.06.2022','institution':'Белорусская государственная орденов Октябрьской Революции и Трудового Красного Знамени сельскохозяйственная академия','speciality':'Сельское строительство и обустройство территорий','qualification':'Инженер','education_level':'Высшее'}],
            'stage_years': '3 года',
            'stage_years_here': '2 года',
            'trudovye_numbers': ['ПК № 2282231'],
            'trudovaya_form2_text': ['Трудовая книжка ПК № 2282231', 'Приказ № 1-к от 16.04.2024 г.'],
            'order_number': '1-к',
            'hire_date': '16.04.2024',
        },
        {
            'fio': 'Горбунов Олег Васильевич',
            'position': 'Заместитель директора-главный инженер',
            'education_full_text': [
                'Высшее',
                'Диплом ВСБ № 0547969 выдан 30.06.2003 г.',
                'Омский государственный технический университет',
                'Экономика и управление на предприятии',
                'Экономист-менеджер',
                'Высшее',
                'Диплом ВСГ 4902664 выдан 30.06.2009 г.',
                'Томский государственный архитектурно-строительный университет',
                'Промышленное и гражданское строительство',
                'Инженер-строитель',
            ],
            'diplomas': [
                {'number':'ВСБ № 0547969','date':'30.06.2003','institution':'Омский государственный технический университет','speciality':'Экономика и управление на предприятии','qualification':'Экономист-менеджер','education_level':'Высшее'},
                {'number':'ВСГ 4902664','date':'30.06.2009','institution':'Томский государственный архитектурно-строительный университет','speciality':'Промышленное и гражданское строительство','qualification':'Инженер-строитель','education_level':'Высшее'},
            ],
            'stage_years': '49 лет',
            'stage_years_here': 'Менее года',
            'trudovye_numbers': ['Трудовая книжка б/н', 'Вкладыш б/н', 'Вкладыш ПК № 00154290', 'Вкладыш АТ-III № 2528394'],
            'trudovaya_form2_text': ['Трудовая книжка б/н', 'Вкладыш б/н', 'Вкладыш ПК № 00154290', 'Вкладыш АТ-III № 2528394', 'Приказ № 3-п от 12.05.2026 г.'],
            'order_number': '3-п',
            'hire_date': '12.05.2026',
            'attestat_number': 'СТ №251921',
            'attestat_form2_text': ['СТ № 251921 26.06.2026 г.', 'Главный инженер (общестроительные работы)'],
            'attestat_form5_text': ['СТ №251921 от 26.06.2026 г. по 26.06.2031 г.', 'Главный инженер (общестроительные работы)'],
        },
        {
            'fio': 'Крот Евгений Васильевич',
            'position': 'Производитель работ (прораб)',
            'education_full_text': [
                'Среднее-специальное',
                'Диплом ЕТ № 113129 выдан 01.03.1986 г.',
                'Витебский индустриальный техникум',
                'Промышленное и гражданское строительство',
                'Техник-строитель',
            ],
            'diplomas': [{'number':'ЕТ № 113129','date':'01.03.1986','institution':'Витебский индустриальный техникум','speciality':'Промышленное и гражданское строительство','qualification':'Техник-строитель','education_level':'Среднее-специальное'}],
            'stage_years': '28 лет',
            'stage_years_here': '1 год',
            'trudovye_numbers': ['ВТ-I № 2206886'],
            'trudovaya_form2_text': ['Трудовая книжка ВТ-I № 2206886', 'Приказ № 2-к от 01.07.2024 г.'],
            'order_number': '2-к',
            'hire_date': '01.07.2024',
            'attestat_number': 'СТ №241674',
            'attestat_form2_text': ['СТ №241674 от 28.11.2025 г.', 'Производитель работ (прораб) (общестроительные работы)'],
            'attestat_form5_text': ['СТ №241674 с 28.11.2025 г. по 28.11.2030 г.', 'Производитель работ (прораб) (общестроительные работы)'],
        },
    ],
}

out = BASE / 'tests' / '_output'
out.mkdir(parents=True, exist_ok=True)
result = generate_company_attestation_package_v2(company, att)
print('warnings:', result['warnings'])
for doc in result['docs']:
    path = out / doc['name']
    path.write_bytes(doc['bytes'])
    print(path)

# Простой текстовый регресс-тест
all_text = {}
for path in sorted(out.glob('*.docx')):
    d = Document(path)
    text = '\n'.join([p.text for p in d.paragraphs] + [c.text for t in d.tables for row in t.rows for c in row.cells])
    all_text[path.name] = text
    assert 'МонТехБел' not in text, path.name
    assert 'Кулешов' not in text, path.name
    assert 'АК СтройФемили' in text, path.name
    assert 'А.Е. Крот' in text, path.name

statement = next(v for k,v in all_text.items() if '1. Заявление' in k)
for token in ['BY94POIS30120162570701933001','ОАО «Паритетбанк»','POISBY2X','7.2.1','7.3.4','7.4.6','7.6.5']:
    assert token in statement, token
forma2 = next(v for k,v in all_text.items() if '2. Форма' in k)
for token in ['Общая численность работающих 19 чел.','Заместитель директора-главный инженер','Производитель работ (прораб)','49 лет','Менее года','Монтажник строительных конструкций','Плотник-бетонщик','Слесарь строительный','Стропальщик','Арматурщик','Электросварщик ручной сварки','Каменщик','Такелажник','Бетонщик','Подсобный рабочий','Маляр','Изолировщик на антикоррозионной изоляции','Кровельщик по металлическим кровлям','Кровельщик по рулонным кровлям и по кровлям из штучных материалов','Штукатур','Землекоп']:
    assert token in forma2, token
forma4 = next(v for k,v in all_text.items() if '4. Форма' in k)
assert 'ВСБ № 0547969' in forma4 and 'ВСГ 4902664' in forma4
forma5 = next(v for k,v in all_text.items() if '5. Форма' in k)
assert 'Горбунов Олег Васильевич' in forma5 and 'Крот Евгений Васильевич' in forma5
print('ALL TEXT ASSERTIONS PASSED')

# Правила автоподстановки и ручных исключений
from generator_company_att import resolve_work_items, resolve_workers, select_relevant_periods
partial_doc = {
    'work_scope_text': 'Общестроительные работы',
    'work_items': ['7.4.1', '7.4.3', '7.4.4'],
    'work_items_source': 'document',
    'workers': [{'profession':'Каменщик','razryad':'IV','count':1,'source':'document'}],
}
resolved_items = resolve_work_items(partial_doc)
assert all(code in resolved_items for code in ['7.2','7.3','7.4','7.5','7.6'])
resolved_workers = resolve_workers(partial_doc, resolved_items)
assert len(resolved_workers) == 16, len(resolved_workers)
assert any(w['profession'] == 'Каменщик' for w in resolved_workers)
partial_doc['excluded_workers'] = ['землекоп|IV']
assert not any(w['profession'] == 'Землекоп' for w in resolve_workers(partial_doc, resolved_items))
assert resolve_work_items({
    'work_scope_text':'Общестроительные работы',
    'work_items':['7.4.1'],
    'work_items_source':'document_exact',
}) == ['7.4.1']

periods = [
    {'start':'01.01.2020','end':'31.12.2020','position':'Бухгалтер'},
    {'start':'01.01.2021','end':'31.12.2022','position':'Мастер'},
    {'start':'01.01.2023','end':'31.12.2024','position':'Производитель работ (прораб)'},
]
selected = select_relevant_periods({'position':'Заместитель директора-главный инженер'}, periods)
assert len(selected) == 2 and all('Бухгалтер' not in p['position'] for p in selected)

index_text = (BASE / 'index.html').read_text('utf-8')
assert "work_items: ['7.2','7.3','7.4','7.5','7.6']" in index_text
assert index_text.count("source:'auto'") >= 3
server_text = (BASE / 'server.py').read_text('utf-8')
assert 'ЗАПОЛНЕННЫЕ ФОРМЫ №1–5 ЯВЛЯЮТСЯ ИСТОЧНИКОМ ИСТИНЫ' in server_text
print('AUTOFILL / EXPERIENCE / FRONTEND ASSERTIONS PASSED')
